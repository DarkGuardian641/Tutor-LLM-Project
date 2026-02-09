import os
from docx import Document
from docx.shared import Inches
import re

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Images: ![caption](path)
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption = match.group(1)
                img_path = match.group(2)
                # Remove file:/// prefix if exists
                img_path = img_path.replace('file:///', '')
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph(caption, style='Caption')
                else:
                    doc.add_paragraph(f"[Image Missing: {caption}]")
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
        
        # Plain text
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Saved to {docx_path}")

if __name__ == "__main__":
    MD_FILE = r"C:\Users\athar\.gemini\antigravity\brain\7fe08c0c-213b-4759-a6cc-bce3f44fb1e2\walkthrough.md"
    DOCX_FILE = r"c:\Visual Studio Code\Tutor LLM Project\data\user_login_implementation.docx"
    
    # Ensure data dir exists
    os.makedirs(os.path.dirname(DOCX_FILE), exist_ok=True)
    
    convert_md_to_docx(MD_FILE, DOCX_FILE)
